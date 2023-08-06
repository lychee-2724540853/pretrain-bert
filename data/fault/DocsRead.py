import os
import random
from docx import Document


class DocxReader():
    """ word文件处理类"""

    def __init__(self, filepath):
        """
        :param filepath: 源文件/路径
        :param destpath: 目标路径
        """
        self.filepath = filepath
        self.docs = []

    def __repr__(self):
        return "word文件读取类"

    def convert2pretrain(self, destpath):
        if not os.path.exists(destpath):
            os.mkdir(destpath)

        if os.path.isfile(self.filepath):
            paras = Document(self.filepath).paragraphs
            destfile = open(os.path.join(destpath, '1.txt'), 'w', encoding='utf8')
            for para in paras:
                destfile.write(para.text+"\n")
            destfile.close()
        else:
            filelist = os.listdir(self.filepath)
            for ids, filename in enumerate(filelist):
                destfile = open(os.path.join(destpath, str(ids)+'.txt'), 'w', encoding='utf8')
                paras = Document(os.path.join(self.filepath, filename)).paragraphs
                for para in paras:
                    destfile.write(para.text+'\n')
                destfile.close()

    def process(self, destpath):
        if os.path.isfile(self.filepath):
            """单个文件"""
            sentences = self.parseDocx(self.filepath)
            nspCorpus = self.nsp(sentences)
            print(nspCorpus)
            print(len(nspCorpus))
        else:
            """ 文件路径 """
            filelist = os.listdir(self.filepath)
            for filename in filelist:
                sentences = self.parseDocx(os.path.join(self.filepath, filename))
                nspCorpus = self.nsp(sentences)
                print(len(nspCorpus))
                print(nspCorpus)

    def parseDocx(self, filepath):
        """ 根据 docx 文件路径，返回文件中的所有句子"""
        doc = Document(filepath)
        sentences = []
        for para in doc.paragraphs:
            texts = para.text.split('。')
            for line in texts:
                line = line.strip()
                if len(line) > 1:
                    sentences.append(line+"。")
        return sentences

    def nsp(self, sentences: list):
        """ 构建 NSP 任务数据集"""
        nspCorpus = []
        for ids, sentence in enumerate(sentences):
            if ids == len(sentences) - 1:
                continue
            ## 正样本
            nspCorpus.append([1, sentence, sentences[ids + 1]])
            first = random.randrange(0, len(sentences))
            while first == ids:
                first = random.randrange(0, len(sentences))
                ## 负样本 1
            nspCorpus.append([0, sentence, sentences[first]])
            second = random.randrange(0, len(sentences))
            while second == first or second == ids:
                second = random.randrange(0, len(sentences))
                ## 负样本 2
            nspCorpus.append([0, sentence, sentences[second]])
        return nspCorpus

if __name__ == '__main__':
    path = 'words/file1.docx'
    desePath = './corpus'
    docReader = DocxReader('words')
    # docReader.process(desePath)
    docReader.convert2pretrain('./txt')
